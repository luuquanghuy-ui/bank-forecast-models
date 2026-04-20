"""
Tạo file Word Chương 3 - Bộ Dữ Liệu, Biến Số và Quy Trình Tiền Xử Lý
(Không liên quan đến LK1)
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_FILE = "d:/labs 2/DOANPTDLKD/thesis_documentation/PART_0_FOUNDATION/CHƯƠNG_3_BaoCao_ChuyenNghiep.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def add_bullet(doc, prefix, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    r1 = p.add_run(prefix)
    r1.bold = True
    r1.font.size = Pt(size)
    r2 = p.add_run(text)
    r2.font.size = Pt(size)
    return p

def add_table(doc, headers, rows, header_color="1B4F72"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_bg(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(9)
    for ri, rd in enumerate(rows):
        rc = table.rows[ri + 1].cells
        bg = "F8F9FA" if ri % 2 == 0 else "FFFFFF"
        for ci, v in enumerate(rd):
            rc[ci].text = str(v)
            set_cell_bg(rc[ci], bg)
            for p in rc[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================
# TITLE
# ============================================================
t = doc.add_heading("CHƯƠNG 3", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = doc.add_heading("BỘ DỮ LIỆU, BIẾN SỐ VÀ QUY TRÌNH TIỀN XỬ LÝ", level=1)
s.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ============================================================
# 3.1 NGUỒN DỮ LIỆU
# ============================================================
add_heading(doc, "3.1. NGUỒN DỮ LIỆU VÀ CẤU TRÚC BỘ DỮ LIỆU", level=1)

add_heading(doc, "3.1.1. Tổng Quan Nguồn Dữ Liệu", level=2)
add_para(doc, "Nghiên cứu thu thập dữ liệu từ ba nguồn độc lập, khoảng thời gian 10 năm (04/02/2016 — 27/02/2026), tương đương 2,495-2,510 ngày giao dịch mỗi mã.")

add_table(doc,
    ["Nguồn", "Thông tin", "Tần suất", "Vai trò"],
    [
        ["Yahoo Finance", "OHLCV 3 mã ngân hàng", "Daily", "Dữ liệu vi mô — hành vi giá cổ phiếu"],
        ["Investing.com", "VNIndex, VN30", "Daily", "Dữ liệu vĩ mô — chỉ số thị trường"],
        ["NHNN VN", "USD/VND, lãi suất liên ngân hàng", "Daily", "Dữ liệu vĩ mô — chi phí vốn"],
    ])

add_heading(doc, "3.1.2. Ba Mã Cổ Phiếu Ngân Hàng Đại Diện", level=2)

add_para(doc, "BID — Ngân hàng TMCP Đầu tư và Phát triển Việt Nam (BIDV)", bold=True)
add_para(doc, "Ngân hàng quốc doanh có quy mô tài sản lớn nhất. Cổ phiếu có mức giá thấp, thanh khoản trung bình, chịu ảnh hưởng đáng kể từ dòng tiền nhà đầu tư cá nhân.")

add_para(doc, "CTG — Ngân hàng TMCP Công thương Việt Nam (VietinBank)", bold=True)
add_para(doc, "Ngân hàng quốc doanh có mức thanh khoản cực kỳ cao. Cổ phiếu CTG rất nhạy cảm với các yếu tố vĩ mô và sự thay đổi chính sách tiền tệ.")

add_para(doc, "VCB — Ngân hàng TMCP Ngoại thương Việt Nam (Vietcombank)", bold=True)
add_para(doc, "Ngân hàng thương mại cổ phần tư nhân lớn nhất (tính theo vốn hóa). Cổ phiếu blue-chip dẫn dắt thị trường, thường biến động sát với chỉ số VNIndex.")

# ============================================================
# 3.2 CẤU TRÚC DATASET
# ============================================================
add_heading(doc, "3.2. CẤU TRÚC BỘ DỮ LIỆU", level=1)

add_para(doc, "Dữ liệu được tiền xử lý và tách thành ba bộ độc lập: banks_BID_dataset.csv, banks_CTG_dataset.csv, banks_VCB_dataset.csv — tránh data leakage giữa các mã. Từ 17 cột gốc, quá trình Feature Engineering tạo ra 25 features có thể sử dụng (bảng dưới).")

add_heading(doc, "3.2.1. 17 Cột Trong Dataset Gốc", level=2)
add_table(doc,
    ["#", "Tên cột", "Loại", "Mô tả", "Nguồn"],
    [
        ["1", "date", "Index", "Ngày giao dịch", "Yahoo Finance"],
        ["2", "open", "Raw", "Giá mở cửa", "Yahoo Finance"],
        ["3", "high", "Raw", "Giá cao nhất ngày", "Yahoo Finance"],
        ["4", "low", "Raw", "Giá thấp nhất ngày", "Yahoo Finance"],
        ["5", "close", "Raw", "Giá đóng cửa điều chỉnh", "Yahoo Finance"],
        ["6", "volume", "Raw", "Khối lượng giao dịch", "Yahoo Finance"],
        ["7", "ticker", "Raw", "Mã cổ phiếu (BID/CTG/VCB)", "—"],
        ["8", "vnindex_close", "Raw", "VNIndex đóng cửa (đã chuẩn hóa)", "Investing.com"],
        ["9", "vn30_close", "Raw", "VN30 đóng cửa (đã chuẩn hóa)", "Investing.com"],
        ["10", "log_return", "Derived", "Lợi suất log: ln(close_t/close_{t-1})", "Derived"],
        ["11", "volatility_20d", "Derived", "Độ biến động 20 ngày", "Derived"],
        ["12", "ma20", "Derived", "Trung bình động 20 ngày", "Derived"],
        ["13", "ma50", "Derived", "Trung bình động 50 ngày", "Derived"],
        ["14", "rsi", "Derived", "RSI chu kỳ 14", "Derived"],
        ["15", "usd_vnd", "Raw", "Tỷ giá USD/VND trung tâm", "NHNN"],
        ["16", "interest_rate", "Raw", "Lãi suất liên ngân hàng (%)", "NHNN"],
        ["17", "index", "Index", "Chỉ số dòng", "—"],
    ])

# ============================================================
# 3.3 FEATURE ENGINEERING
# ============================================================
add_heading(doc, "3.3. FEATURE ENGINEERING — 25 FEATURES", level=1)
add_para(doc, "Từ 17 cột gốc, Feature Engineering tạo ra các biến Lagged và Ratio. Tổng cộng 25 features có thể sử dụng cho các mô hình dự báo.")

add_heading(doc, "3.3.1. Technical Indicators (12 features)", level=2)
add_table(doc,
    ["STT", "Tên biến", "Loại", "Công thức / Mô tả"],
    [
        ["1", "open", "Raw", "Giá mở cửa ngày giao dịch"],
        ["2", "high", "Raw", "Giá cao nhất ngày"],
        ["3", "low", "Raw", "Giá thấp nhất ngày"],
        ["4", "close", "Raw", "Giá đóng cửa ngày"],
        ["5", "volume", "Raw", "Khối lượng giao dịch"],
        ["6", "log_return", "Derived", "ln(close_t / close_{t-1})"],
        ["7", "volatility_20d", "Derived", "std(log_return_{t-19:t})"],
        ["8", "ma20", "Derived", "Trung bình close 20 ngày"],
        ["9", "ma50", "Derived", "Trung bình close 50 ngày"],
        ["10", "rsi", "Derived", "RSI chu kỳ 14"],
        ["11", "volume_ratio", "Derived", "volume / rolling(20).mean()"],
        ["12", "ma_ratio", "Derived", "close / ma20"],
    ])

add_heading(doc, "3.3.2. Lagged & Macro Features (13 features)", level=2)
add_table(doc,
    ["STT", "Tên biến", "Loại", "Công thức / Mô tả"],
    [
        ["13", "return_lag1", "Lag", "log_return.shift(1)"],
        ["14", "return_lag2", "Lag", "log_return.shift(2)"],
        ["15", "return_lag3", "Lag", "log_return.shift(3)"],
        ["16", "return_lag5", "Lag", "log_return.shift(5)"],
        ["17", "volatility_lag1", "Lag", "volatility_20d.shift(1)"],
        ["18", "volatility_lag2", "Lag", "volatility_20d.shift(2)"],
        ["19", "rsi_lag1", "Lag", "rsi.shift(1)"],
        ["20", "vnindex_close", "Macro", "VNIndex đóng cửa ngày"],
        ["21", "vn30_close", "Macro", "VN30 đóng cửa ngày"],
        ["22", "usd_vnd", "Macro", "Tỷ giá USD/VND"],
        ["23", "interest_rate", "Macro", "Lãi suất liên ngân hàng (%)"],
        ["24", "vnindex_lag1", "Macro Lag", "vnindex_close.shift(1)"],
        ["25", "vn30_lag1", "Macro Lag", "vn30_close.shift(1)"],
    ])

# ============================================================
# 3.4 TIỀN XỬ LÝ
# ============================================================
add_heading(doc, "3.4. TIỀN XỬ LÝ DỮ LIỆU", level=1)

add_heading(doc, "3.4.1. Xử Lý Giá Trị Thiếu", level=2)
add_para(doc, "Dữ liệu tài chính (OHLCV) có độ toàn vẹn 100%. Tuy nhiên, khi trích xuất các biến Lagged và MA, NaN xuất hiện tự nhiên tại các dòng đầu tiên (ví dụ: ma50 cần 50 ngày).")
add_bullet(doc, "Quyết định: ", "Drop ~50 dòng đầu tiên (~2% tổng số dòng) thay vì Forward Fill. Forward Fill tạo 'bước nhảy phẳng' làm sai lệch tín hiệu.")

add_heading(doc, "3.4.2. Xử Lý Ngày Nghỉ Lễ", level=2)
add_para(doc, "Thị trường Việt Nam có kỳ nghỉ lễ Âm lịch kéo dài (7-10 ngày).")
add_bullet(doc, "Không Forward Fill: ", "Nếu điền giá ngày cận Tết cho cả kỳ nghỉ, log_return bị ép bằng 0. Không có dữ liệu không có nghĩa là giá không đổi.")
add_bullet(doc, "Giải pháp: ", "Giữ Gap thời gian. Log_return tính trên ngày giao dịch thực tế gần nhất (Gap-aware).")

add_heading(doc, "3.4.3. Xử Lý Giá Trị Ngoại Lai", level=2)
add_para(doc, "Theo phân phối chuẩn, Outliers (> 3 độ lệch chuẩn) chỉ ~0.27%. Nhưng log_return có tới 3%-5% Outliers — bằng chứng của Fat Tails đặc thù chứng khoán.")
add_bullet(doc, "Quyết định: ", "Giữ nguyên 100% Outliers. Đây không phải sai số kỹ thuật mà là 'tín hiệu vàng' từ COVID-19 và các cú sốc thị trường, giúp thuật toán học cách phản ứng với khủng hoảng.")

# ============================================================
# 3.5 STATIONARITY
# ============================================================
add_heading(doc, "3.5. KIỂM ĐỊNH TÍNH DỪNG (ADF TEST)", level=1)
add_para(doc, "Nghiên cứu sử dụng Augmented Dickey-Fuller (ADF) Test. H0: chuỗi có nghiệm đơn vị (không dừng). H1: chuỗi dừng (p-value < 0.05).")

add_table(doc,
    ["Biến", "BID p-value", "CTG p-value", "VCB p-value", "Kết luận"],
    [
        ["Giá đóng cửa (close)", "0.7978", "0.9781", "0.6949", "Không Dừng"],
        ["Lợi suất Log (log_return)", "0.0000", "0.0000", "0.0000", "Dừng ✓"],
        ["Biến động (volatility_20d)", "0.0005", "0.0001", "0.0001", "Dừng ✓"],
    ])

add_para(doc, "close: Không dừng — mức giá thô mang xu hướng tích luỹ (Martingale). Đưa trực tiếp vào mô hình gây Spurious correlation.", bold=True)
add_para(doc, "log_return: Rất dừng (p < 0.0001) — đã loại bỏ lạm phát giá. Biến phù hợp nhất cho dự báo.", bold=True)
add_para(doc, "volatility_20d: Dừng — có tính mean-reversion, không vỡ vụn. Phù hợp cho GARCH.", bold=True)

# ============================================================
# 3.6 EDA
# ============================================================
add_heading(doc, "3.6. PHÂN TÍCH THỐNG KÊ MÔ TẢ (EDA)", level=1)
add_para(doc, "Phân phối lợi suất đạt chuẩn khi Skewness = 0 và Excess Kurtosis = 0. Thực tế hoàn toàn khác — bằng chứng của Fat Tails và Volatility Clustering.")

add_heading(doc, "3.6.1. Thống Kê Mô Tả", level=2)
add_table(doc,
    ["Đặc trưng", "BID", "CTG", "VCB", "Giải thích"],
    [
        ["Skewness (return)", "-0.121", "-0.147", "-0.008", "Âm — đuôi trái dài, rơi nhanh hơn lên"],
        ["Kurtosis (return)", "2.208", "2.297", "3.090", "Dương — Fat tails, xác suất cực đoan cao"],
        ["Kurtosis (Volume)", "29.25", "4.435", "48.15", "Cực đoan — volume đôi khi giật gấp 5-10 lần TB"],
    ])

add_heading(doc, "3.6.2. Risk DNA — Đặc Điểm Từng Mã", level=2)

add_para(doc, "BID — Ngân hàng Momentum Bán lẻ", bold=True)
add_para(doc, "• Kurtosis Volume = 29.25 — giật cục cực cao")
add_para(doc, "• VN30 rank #13 — ít đi theo thị trường chung")
add_para(doc, "• Profile: Cổ phiếu bị chi phối bởi FOMO nhà đầu tư cá nhân")

add_para(doc, "CTG — Ngân hàng Nhạy Cảm Vĩ Mô", bold=True)
add_para(doc, "• Skewness = -0.147 — lệch âm mạnh nhất, rơi nhanh hơn lên")
add_para(doc, "• USD/VND nhạy cảm nhất trong 3 banks")
add_para(doc, "• Profile: Ngân hàng chịu ảnh hưởng mạnh bởi yếu tố vĩ mô")

add_para(doc, "VCB — Ngân hàng Blue-Chip Dẫn Dắt", bold=True)
add_para(doc, "• VN30 rank #4 — đi theo thị trường chung mạnh nhất")
add_para(doc, "• Kurtosis Volume = 48.15 — cực đoan nhất")
add_para(doc, "• Profile: Blue-chip hàng đầu VN, bị ảnh hưởng bởi dòng tiền tổ chức")

# ============================================================
# 3.7 TÓM TẮT
# ============================================================
add_heading(doc, "3.7. TÓM TẮT CHƯƠNG 3", level=1)
add_table(doc,
    ["Hạng mục", "Chi tiết"],
    [
        ["Dataset", "3 file độc lập: BID, CTG, VCB"],
        ["Thời gian", "04/02/2016 — 27/02/2026 (~2,500 ngày/mã)"],
        ["Cột gốc", "17 cột (OHLCV, chỉ số, chỉ báo, macro)"],
        ["Features sau FE", "25 features (lag, ratio, derived)"],
        ["Missing Values", "Drop ~50 dòng đầu (~2%)"],
        ["Holiday", "Giữ Gap — không Forward Fill"],
        ["Outliers", "Giữ nguyên 100% — tín hiệu thật"],
        ["ADF Test", "log_return & volatility_20d: Dừng ✓; close: Không Dừng"],
        ["Key Finding", "Fat tails (Kurtosis >> 0), Volatility clustering"],
    ])

doc.add_paragraph()
add_para(doc, "Chương 3 đã tái thiết lập bộ dữ liệu hoàn chỉnh: 17 cột gốc, 25 features, xử lý Missing Values/Outliers, khẳng định tính Dừng của dữ liệu lợi suất, và phân tích Fat tails. Móng đã xây vững chãi chờ các mô hình dự báo ở Chương 4.")

doc.save(OUTPUT_FILE)
print(f"Saved: {OUTPUT_FILE}")
